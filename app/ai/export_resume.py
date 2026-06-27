from docx import Document


def export_resume_to_docx(

    tailored_bullets,

    output_file
):

    document = Document()

    document.add_heading(

        "Tailored Resume",

        level=1
    )

    document.add_heading(

        "Professional Experience",

        level=2
    )

    for bullet in tailored_bullets:

        document.add_paragraph(

            bullet,

            style="List Bullet"
        )

    document.save(output_file)

    print(

        f"\nResume exported to: {output_file}"
    )